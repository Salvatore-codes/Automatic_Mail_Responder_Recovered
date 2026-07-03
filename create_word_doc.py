import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell margins (padding) in twentieths of a point (dxa)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    """Set background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_borders(cell, color="CCCCCC", sz="4", val="single"):
    """Set borders for a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def create_document():
    doc = Document()
    
    csv_path = os.path.join(os.path.dirname(__file__), "data", "sku_catalog.csv")
    
    # Page Margins

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette Definitions
    COLOR_PRIMARY = RGBColor(11, 15, 25)      # Slate / Dark Navy
    COLOR_INDIGO = RGBColor(79, 70, 229)     # Deep Indigo
    COLOR_MUTED = RGBColor(100, 116, 139)    # Gray Muted
    COLOR_GREEN = RGBColor(16, 185, 129)     # Success Green
    
    # Document Style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(30, 41, 59) # Off-black text

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("PROPOSAL & SPECS")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_INDIGO
    
    title_h = doc.add_paragraph()
    title_h.paragraph_format.space_after = Pt(12)
    run_h = title_h.add_run("Automated SKU Matching & Quotation Engine")
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(24)
    run_h.font.bold = True
    run_h.font.color.rgb = COLOR_PRIMARY

    # Subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run("A comparative analysis and roadmap for local free pipelines vs. paid API integrations.")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_MUTED
    
    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(24)
    p_div_run = p_div.add_run("—" * 60)
    p_div_run.font.color.rgb = COLOR_MUTED
    
    # 1. Executive Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run_h1 = h1.add_run("1. Executive Summary")
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY
    
    p_exec = doc.add_paragraph(
        "A typical hardware store processes hundreds of orders daily via WhatsApp, email, and paper scans. "
        "With a catalog of 10,000+ SKUs, mapping colloquial or misspelled text into exact database IDs to draft quotations "
        "is a massive manual bottleneck. This proposal evaluates two technical approaches to automate this workflow: "
        "a zero-cost local fuzzy logic model (Scenario A) and an intelligent, semantic API-driven hybrid pipeline (Scenario B)."
    )
    p_exec.paragraph_format.space_after = Pt(12)
    
    # 2. Core Concepts
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    run_h2 = h2.add_run("2. The Matching Methodologies")
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = COLOR_PRIMARY

    # Table comparing Scenario A vs B
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    # Set headers
    headers = ["Metric / Feature", "Scenario A: Free & Local", "Scenario B: Paid API-Driven"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F2937")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        
    data = [
        ("Information OCR Ingestion", "Basic (Tesseract OCR). Struggling with tilted text, tables, or hand-written lists.", "Exceptional (Gemini Vision). Easily transcribes paper lists, tables, and handwriting."),
        ("Parsing & Itemization", "Heuristics (Regex). Fails on compound sentences (e.g. \"50 bolts and matching nuts\" gets matched as one item).", "Semantic Context (LLM). Understands multi-product lines and splits them into distinct items."),
        ("Synonym Matching", "Lexical (RapidFuzz / spelling matching). Misses context entirely if terms look different (e.g. \"teflon taps\" vs \"teflon tape\").", "Vector Search (Embeddings). Translates colloquial descriptions into database IDs based on intent/meaning."),
        ("Hosting & Monthly Fees", "Free ($0.00). Runs completely on standard local store PC.", "Pay-per-use (Est. $5-$15/month). Fractions of a cent per API request."),
        ("Development Cost", "High maintenance. Requires constantly updating code and regex rules for typos.", "Low maintenance. Large Language Models adapt dynamically without code changes.")
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            cells[col_idx].text = text
            set_cell_margins(cells[col_idx], top=100, bottom=100)
            if col_idx == 0:
                cells[col_idx].paragraphs[0].runs[0].font.bold = True
            elif col_idx == 1:
                set_cell_background(cells[col_idx], "F8FAFC")
            elif col_idx == 2:
                set_cell_background(cells[col_idx], "EEF2F6")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. Production Roadmap
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)
    run_h3 = h3.add_run("3. Critical Production integrations")
    run_h3.font.size = Pt(16)
    run_h3.font.bold = True
    run_h3.font.color.rgb = COLOR_PRIMARY
    
    p_road = doc.add_paragraph(
        "To implement this successfully, the system must bridge the gap between simple text matching "
        "and active shop workflows. The following five modules have been successfully simulated in our prototype:"
    )
    p_road.paragraph_format.space_after = Pt(12)
    
    bullet_items = [
        ("CRM Profile Lookup", "Querying customer databases using incoming email/phone records to dynamically apply special discount groups (e.g. Contractors get 15% wholesale discounts)."),
        ("ERP Stock Level Check", "Synchronizing with POS/inventory databases to query availability in real-time and alert the operator about out-of-stock items directly on the draft invoice."),
        ("Human-in-the-Loop (HITL) Gate", "An automated quality control gate. If a matched SKU holds a confidence level below 80%, the order pauses and lists the top 3 alternative catalog SKUs for quick manual confirmation."),
        ("Feedback Learning Loop", "Capturing operator corrections and manual SKU selections. Overrides are written directly to a local synonyms file so that future encounters automatically match with 100% confidence."),
        ("Trigger-driven Ingestion", "Automated connectors that watch and ingest orders directly from WhatsApp Webhooks or Gmail Pub/Sub listeners.")
    ]
    
    for title, desc in bullet_items:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(4)
        run_bold = p_bullet.add_run(f"{title}: ")
        run_bold.font.bold = True
        p_bullet.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. Running the Prototype
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(8)
    run_h4 = h4.add_run("4. Running the Prototype Simulator")
    run_h4.font.size = Pt(16)
    run_h4.font.bold = True
    run_h4.font.color.rgb = COLOR_PRIMARY
    
    p_run = doc.add_paragraph(
        "The Python code folder contains a CLI simulator matching these production specs. "
        "You can run the demo from your terminal inside the project directory:"
    )
    p_run.paragraph_format.space_after = Pt(8)
    
    # Code block container
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.5)
    p_code.paragraph_format.space_after = Pt(12)
    run_code = p_code.add_run(
        "cd D:\\sku-matcher-prototype\n"
        "& \"C:\\Users\\Admin\\AppData\\Local\\Programs\\Python\\Python314\\python.exe\" run_demo.py"
    )

    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(9.5)
    run_code.font.color.rgb = COLOR_PRIMARY
    
    # 5. Recommended Tools & Technology Stack
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(8)
    run_h5 = h5.add_run("5. Recommended Tools & Technology Stack")
    run_h5.font.size = Pt(16)
    run_h5.font.bold = True
    run_h5.font.color.rgb = COLOR_PRIMARY
    
    p_tools = doc.add_paragraph(
        "The following matrix outlines the recommended software tools and technology stacks for "
        "each pipeline stage in a production environment:"
    )
    p_tools.paragraph_format.space_after = Pt(12)
    
    # Tools Table
    table_tools = doc.add_table(rows=8, cols=4)
    table_tools.style = 'Table Grid'
    
    tool_headers = ["Pipeline Stage", "Scenario A: Free Stack", "Scenario B: Paid/Hybrid", "Purpose / Function"]
    hdr_cells_t = table_tools.rows[0].cells
    for i, title in enumerate(tool_headers):
        hdr_cells_t[i].text = title
        set_cell_background(hdr_cells_t[i], "1F2937")
        run = hdr_cells_t[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(hdr_cells_t[i], top=120, bottom=120)
        
    tool_data = [
        ("Ingestion & Triggers", "Python + Local folder listener / IMAP", "n8n / Make.com + Twilio / Gmail API", "Triggers quotation runs on email or chat receipt."),
        ("OCR (Scanned Images)", "Tesseract OCR / EasyOCR", "Gemini 2.5 Flash API (Multimodal)", "Converts photos of handwritten or printed orders to text."),
        ("Structured Parser", "Python Regular Expressions (re)", "Gemini 2.5 Flash (Structured Outputs)", "Extracts items, specs, and quantities into JSON format."),
        ("SKU Matching", "RapidFuzz + Custom TF-IDF similarity", "Gemini Embeddings + PostgreSQL (pgvector)", "Performs semantic search against 10,000+ stock IDs."),
        ("CRM & ERP Check", "Local CSV Files / SQLite DB", "Odoo / QuickBooks / SAP REST APIs", "Validates stock levels and applies dealer discounts."),
        ("Review Dashboard", "CLI Terminal Prompts", "Retool / Glide / Custom React GUI", "Human-in-the-Loop interface for manual overrides."),
        ("Quote Dispatch", "Python standard smtplib (SMTP mail)", "SendGrid / Twilio WhatsApp Business API", "Auto-sends final quotation PDFs back to enquiries.")
    ]
    
    for row_idx, row_data in enumerate(tool_data, start=1):
        cells = table_tools.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            cells[col_idx].text = text
            set_cell_margins(cells[col_idx], top=100, bottom=100)
            if col_idx == 0:
                cells[col_idx].paragraphs[0].runs[0].font.bold = True
            elif col_idx == 1:
                set_cell_background(cells[col_idx], "F8FAFC")
            elif col_idx == 2:
                set_cell_background(cells[col_idx], "EEF2F6")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    output_path = os.path.join(os.path.dirname(csv_path), "sku_matching_proposal.docx")
    doc.save(output_path)
    print(f"[System] Word Document successfully saved to {output_path}")

if __name__ == "__main__":
    create_document()
